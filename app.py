from fastapi import FastAPI, UploadFile, File, Form, WebSocket, BackgroundTasks, Depends, HTTPException, status, Request
from fastapi.middleware.cors import CORSMiddleware
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
from pypdf import PdfReader
from deep_translator import GoogleTranslator
from fpdf import FPDF
from fpdf.errors import FPDFUnicodeEncodingException
import fitz 

import tempfile
import boto3
import uuid
import asyncio
import os
import re

SCRIPT_FONTS: dict = {}

_FONT_FILES = {
    "Devanagari": "NotoSansDevanagari-Regular.ttf",  # hi, mr
    "Bengali":    "NotoSansBengali-Regular.ttf",      # bn
    "Gujarati":   "NotoSansGujarati-Regular.ttf",     # gu
    "Gurmukhi":   "NotoSansGurmukhi-Regular.ttf",     # pa
    "Tamil":      "NotoSansTamil-Regular.ttf",        # ta
    "Telugu":     "NotoSansTelugu-Regular.ttf",       # te
    "Kannada":    "NotoSansKannada-Regular.ttf",      # kn
    "Malayalam":  "NotoSansMalayalam-Regular.ttf",    # ml
    "Arabic":     "NotoSansArabic-Regular.ttf",       # ar, ur
    "Cyrillic":   "NotoSans-Regular.ttf",             # ru
    "Regular":    "NotoSans-Regular.ttf",             # Generic Unicode fallback
}

def _register_fonts_startup():
    for script, filename in _FONT_FILES.items():
        if os.path.exists(filename):
            SCRIPT_FONTS[script] = filename
            print(f"[fonts] OK  {script} ({filename})")
        else:
            print(f"[fonts] MISSING {filename} — {script} will use fallback")
    _arial = "/Library/Fonts/Arial Unicode.ttf"
    if not os.path.exists(_arial):
        _arial = "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"
    if os.path.exists(_arial):
        for script in ("CJK_SC", "CJK_JP", "CJK_KR"):
            SCRIPT_FONTS[script] = _arial
        print(f"[fonts] OK  CJK_SC/JP/KR (Arial Unicode: {_arial})")
    else:
        for script in ("CJK_SC", "CJK_JP", "CJK_KR"):
            SCRIPT_FONTS[script] = "cjk"
        print(f"[fonts] OK  CJK (using fitz builtin 'cjk')")

_register_fonts_startup()


def _font_for(ch: str, fallback_font: str = "helv") -> str:
    cp = ord(ch)
    if 0x0900 <= cp <= 0x097F or 0xA8E0 <= cp <= 0xA8FF:
        return SCRIPT_FONTS.get("Devanagari", fallback_font)
    if 0x0980 <= cp <= 0x09FF:
        return SCRIPT_FONTS.get("Bengali", fallback_font)
    if 0x0A00 <= cp <= 0x0A7F:
        return SCRIPT_FONTS.get("Gurmukhi", fallback_font)
    if 0x0A80 <= cp <= 0x0AFF:
        return SCRIPT_FONTS.get("Gujarati", fallback_font)
    if 0x0B80 <= cp <= 0x0BFF:
        return SCRIPT_FONTS.get("Tamil", fallback_font)
    if 0x0C00 <= cp <= 0x0C7F:
        return SCRIPT_FONTS.get("Telugu", fallback_font)
    if 0x0C80 <= cp <= 0x0CFF:
        return SCRIPT_FONTS.get("Kannada", fallback_font)
    if 0x0D00 <= cp <= 0x0D7F:
        return SCRIPT_FONTS.get("Malayalam", fallback_font)
    if (0x0600 <= cp <= 0x06FF or 0x0750 <= cp <= 0x077F or
            0xFB50 <= cp <= 0xFDFF or 0xFE70 <= cp <= 0xFEFF):
        return SCRIPT_FONTS.get("Arabic", fallback_font)
    if 0xAC00 <= cp <= 0xD7AF or 0x1100 <= cp <= 0x11FF:
        return SCRIPT_FONTS.get("CJK_KR", fallback_font)
    if 0x3040 <= cp <= 0x309F or 0x30A0 <= cp <= 0x30FF:
        return SCRIPT_FONTS.get("CJK_JP", fallback_font)
    if (0x4E00 <= cp <= 0x9FFF or 0x3400 <= cp <= 0x4DBF or
            0xF900 <= cp <= 0xFAFF):
        return SCRIPT_FONTS.get("CJK_SC", fallback_font)
    if 0x3000 <= cp <= 0x303F or 0x3200 <= cp <= 0x33FF:
        return SCRIPT_FONTS.get("CJK_SC", SCRIPT_FONTS.get("CJK_JP", fallback_font))
    if 0xFF00 <= cp <= 0xFFEF:
        return SCRIPT_FONTS.get("CJK_SC", SCRIPT_FONTS.get("CJK_JP", fallback_font))
    if 0x0400 <= cp <= 0x04FF:
        return SCRIPT_FONTS.get("Cyrillic", fallback_font)
    # --- ADDED: General Punctuation (includes en dash – and em dash —) ---
    if 0x2000 <= cp <= 0x206F:
        return SCRIPT_FONTS.get("Regular", SCRIPT_FONTS.get("Cyrillic", fallback_font))
    if cp > 127:
        # Try to find any loaded Unicode font for symbols/punctuation
        for k in ["Regular", "Cyrillic", "CJK_SC", "ArialUnicode"]:
            if k in SCRIPT_FONTS:
                return SCRIPT_FONTS[k]
    return fallback_font


def draw_mixed_text(pdf_out, text: str, x: float, y: float, font_size: float,
                    fallback_font: str = "helv", lang_code: str = "en"):
    """Draw text switching fonts automatically per Unicode block using fpdf2."""
    if not text:
        return

    segments, cur_font, cur_chunk = [], None, ""
    for ch in text:
        f = _font_for(ch)
        
        if ch.isspace() and cur_font is not None:
            f = cur_font
        if f != cur_font:
            if cur_chunk:
                segments.append((cur_font, cur_chunk))
            cur_font, cur_chunk = f, ch
        else:
            cur_chunk += ch
    if cur_chunk:
        segments.append((cur_font, cur_chunk))

    
    is_deva = any(0x0900 <= ord(c) <= 0x097F for c in text)

    cursor_x = x
    for seg_font, seg_text in segments:
        if seg_font and seg_font.endswith(".ttf"):
            font_id = os.path.basename(seg_font).split(".")[0]
            try:
                if font_id.lower() not in pdf_out.fonts:
                    pdf_out.add_font(font_id, fname=seg_font)
                pdf_out.set_font(font_id, size=font_size)
                
                
                if "Devanagari" in seg_font or is_deva:
                    _hb_lang = {"hi": "hin", "mr": "mar", "ne": "nep", "kok": "kok"}.get(lang_code, "hin")
                    pdf_out.set_text_shaping(True, script="dev2", language=_hb_lang)
                elif "Bengali" in (seg_font or ""):
                    pdf_out.set_text_shaping(True, script="bng2", language="ben")
                elif "Tamil" in (seg_font or ""):
                    pdf_out.set_text_shaping(True, script="tml2", language="tam")
                elif "Telugu" in (seg_font or ""):
                    pdf_out.set_text_shaping(True, script="tel2", language="tel")
                elif "Kannada" in (seg_font or ""):
                    pdf_out.set_text_shaping(True, script="knd2", language="kan")
                elif "Malayalam" in (seg_font or ""):
                    pdf_out.set_text_shaping(True, script="mlm2", language="mal")
                elif "Gujarati" in (seg_font or ""):
                    pdf_out.set_text_shaping(True, script="gjr2", language="guj")
                elif "Gurmukhi" in (seg_font or ""):
                    pdf_out.set_text_shaping(True, script="guru", language="pan")
                else:
                    pdf_out.set_text_shaping(True)  # generic
            except Exception:
                pdf_out.set_font("helvetica", size=font_size)
        else:
            pdf_out.set_font("helvetica", size=font_size)
            pdf_out.set_text_shaping(False) 
        
        pdf_out.text(x=cursor_x, y=y, text=seg_text)
        cursor_x += pdf_out.get_string_width(seg_text)

from sqlalchemy.orm import Session
from database import engine, Base, get_db
import models
import auth
import email_utils

from sqlalchemy import text


try:
    models.Base.metadata.create_all(bind=engine)
    
  
    with engine.begin() as conn:
        try:
            conn.execute(text("ALTER TABLE users ADD COLUMN email VARCHAR UNIQUE"))
        except Exception:
            pass
        try:
            conn.execute(text("ALTER TABLE users ADD COLUMN reset_token VARCHAR"))
        except Exception:
            pass
        try:
            conn.execute(text("ALTER TABLE users ADD COLUMN reset_token_expiry TIMESTAMP"))
        except Exception:
            pass
except Exception as e:
    print(f"Warning: Failed to create/upgrade database tables. Error: {e}")

app = FastAPI()


from fastapi.responses import JSONResponse
import traceback

@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    error_msg = f"Internal Server Error: {str(exc)}"
    print(f"Global Error: {error_msg}")
    traceback.print_exc()
    return JSONResponse(
        status_code=500,
        content={"detail": error_msg}
    )

from sqlalchemy import text

@app.get("/health")
def health_check(db: Session = Depends(get_db)):
    try:
        
        db.execute(text("SELECT 1"))
        return {"status": "healthy", "database": "connected"}
    except Exception as e:
        return JSONResponse(
            status_code=503,
            content={"status": "unhealthy", "database": "disconnected", "error": str(e)}
        )


limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)


progress_store = {}


s3 = boto3.client("s3", region_name="eu-north-1")
BUCKET_NAME = "pdf-translator-storage"


_allowed_origins = os.environ.get("ALLOWED_ORIGIN", "*").split(",")
_allowed_origins = [o.strip() for o in _allowed_origins if o.strip()]

app.add_middleware(
    CORSMiddleware,
    allow_origins=_allowed_origins,
    allow_credentials=True if "*" not in _allowed_origins else False,
    allow_methods=["*"],
    allow_headers=["*"],
)



from fastapi import WebSocketDisconnect

@app.websocket("/progress/{task_id}")
async def progress_socket(websocket: WebSocket, task_id: str):

    await websocket.accept()

    try:
        while True:

            progress = progress_store.get(task_id, 0)

            await websocket.send_json({"progress": progress})

            if progress >= 100 or progress < 0:
                break

            await asyncio.sleep(0.5)
    except WebSocketDisconnect:
        print(f"WebSocket disconnected for task {task_id}")
    except Exception as e:
        print(f"WebSocket error for task {task_id}: {e}")

# ---------------------------
# Auth API Endpoints
# ---------------------------
from fastapi.security import OAuth2PasswordRequestForm

from pydantic import BaseModel, EmailStr

class UserCreate(BaseModel):
    username: str
    email: str
    password: str

@app.post("/register")
def register(user: UserCreate, db: Session = Depends(get_db)):
    try:
        db_user = db.query(models.User).filter(
            (models.User.username == user.username) | (models.User.email == user.email)
        ).first()
        if db_user:
            raise HTTPException(status_code=400, detail="Username or email already registered")
        
        hashed_password = auth.get_password_hash(user.password)
        new_user = models.User(username=user.username, email=user.email, password_hash=hashed_password)
        db.add(new_user)
        db.commit()
        return {"message": "User created successfully"}
    except HTTPException:
        raise
    except Exception as e:
        print(f"Registration Error: {e}")
        db.rollback()
        raise HTTPException(status_code=500, detail="An error occurred during registration. Please try again.")

@app.post("/login")
@limiter.limit("10/minute")
def login(request: Request, form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(get_db)):
    user = db.query(models.User).filter(models.User.username == form_data.username).first()
    if not user or not auth.verify_password(form_data.password, user.password_hash):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token_expires = auth.timedelta(minutes=auth.ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = auth.create_access_token(
        data={"sub": user.username}, expires_delta=access_token_expires
    )
    return {"access_token": access_token, "token_type": "bearer"}

from datetime import datetime

class ForgotPasswordRequest(BaseModel):
    email: str

@app.post("/forgot-password")
@limiter.limit("5/minute")
def forgot_password(request: Request, req: ForgotPasswordRequest, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    user = db.query(models.User).filter(models.User.email == req.email).first()
    if not user:
        return {"message": "If that email exists, an OTP has been sent to it."}
    
    
    import random
    otp = str(random.randint(100000, 999999))
    
    user.reset_token = otp
    user.reset_token_expiry = datetime.utcnow() + auth.timedelta(minutes=10)  
    db.commit()

    
    background_tasks.add_task(email_utils.send_reset_password_email, req.email, otp)

    return {"message": "A 6-digit OTP has been sent to your email. It expires in 10 minutes."}

class ResetPasswordRequest(BaseModel):
    email: EmailStr
    reset_token: str
    new_password: str

@app.post("/reset-password")
def reset_password(request: ResetPasswordRequest, db: Session = Depends(get_db)):
    user = db.query(models.User).filter(
        models.User.email == request.email,
        models.User.reset_token == request.reset_token
    ).first()
    
    if not user or user.reset_token_expiry < datetime.utcnow():
        raise HTTPException(status_code=400, detail="Invalid or expired reset token")
    
    # Update password and clear token
    user.password_hash = auth.get_password_hash(request.new_password)
    user.reset_token = None
    user.reset_token_expiry = None
    db.commit()
    
    return {"message": "Password has been successfully reset. You can now log in."}

class ForgotUsernameRequest(BaseModel):
    email: str

@app.post("/forgot-username")
@limiter.limit("5/minute")
def forgot_username(request: Request, req: ForgotUsernameRequest, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    user = db.query(models.User).filter(models.User.email == req.email).first()
    if not user:
        return {"message": "If that email is registered, your username has been sent to it."}

    background_tasks.add_task(email_utils.send_forgot_username_email, req.email, user.username)

    return {"message": "If that email is registered, your username has been sent to it."}





class SendPdfRequest(BaseModel):
    recipient_email: EmailStr
    download_url: str
    filename: str

@app.post("/send-pdf")
@limiter.limit("5/minute")
def send_pdf(request: Request, req: SendPdfRequest, background_tasks: BackgroundTasks,
             current_user: models.User = Depends(auth.get_current_user)):
    """Fetch the translated PDF from S3 and email it as an attachment to any recipient."""
    def _do_send():
        try:
            import requests as _req
            resp = _req.get(req.download_url, timeout=30)
            resp.raise_for_status()
            pdf_bytes = resp.content
            email_utils.send_pdf_email(
                recipient_email=req.recipient_email,
                filename=req.filename,
                pdf_bytes=pdf_bytes,
            )
        except Exception as e:
            print(f"[send-pdf] Error: {e}")

    background_tasks.add_task(_do_send)
    return {"message": f"PDF is being sent to {req.recipient_email}. It should arrive in a few seconds!"}




@app.get("/history")
def get_history(current_user: models.User = Depends(auth.get_current_user), db: Session = Depends(get_db)):
    tasks = db.query(models.TranslationTask).filter(models.TranslationTask.user_id == current_user.id).order_by(models.TranslationTask.created_at.desc()).all()
    return tasks



from fastapi import Header

def verify_admin(x_admin_secret: str = Header(...)):
    import os
    from dotenv import load_dotenv
    load_dotenv(override=True)
    real_secret = os.environ.get("ADMIN_SECRET")
    if not real_secret or x_admin_secret != real_secret:
        raise HTTPException(status_code=403, detail="Invalid admin secret")
    return True

@app.get("/admin/stats")
def get_admin_stats(db: Session = Depends(get_db), _: bool = Depends(verify_admin)):
    total_users = db.query(models.User).count()
    total_tasks = db.query(models.TranslationTask).count()
    completed_tasks = db.query(models.TranslationTask).filter(models.TranslationTask.status == "completed").count()
    failed_tasks = db.query(models.TranslationTask).filter(models.TranslationTask.status == "failed").count()
    processing_tasks = db.query(models.TranslationTask).filter(models.TranslationTask.status == "processing").count()
    return {
        "users": total_users,
        "translations_total": total_tasks,
        "translations_completed": completed_tasks,
        "translations_failed": failed_tasks,
        "translations_processing": processing_tasks
    }

@app.get("/admin/users")
def get_admin_users(db: Session = Depends(get_db), _: bool = Depends(verify_admin)):
    from sqlalchemy import func
    
    users = db.query(
        models.User.id,
        models.User.username,
        models.User.email,
        func.count(models.TranslationTask.id).label('task_count')
    ).outerjoin(
        models.TranslationTask, models.User.id == models.TranslationTask.user_id
    ).group_by(models.User.id).all()
    
    return [
        {
            "id": u.id,
            "username": u.username,
            "email": u.email,
            "task_count": u.task_count
        } for u in users
    ]



def _translate_with_retry(text: str, target_language: str, max_retries: int = 2) -> str:
    """Translate text to target_language with retry logic.
    Returns the translated string, or the original text on failure."""
    import time
    if not text or not text.strip():
        return text
   
    text = " ".join(text.split())
    if len(text) < 2: return text

    for attempt in range(max_retries + 1):
        try:
            result = GoogleTranslator(source="auto", target=target_language).translate(text)
            if result and result.strip():
                
                return " ".join(result.split())
        except Exception as e:
            print(f"[translate] attempt {attempt + 1} failed: {e}")
            if attempt < max_retries:
                time.sleep(0.5)
    return text  # fall back to original on total failure


def _deduplicate_lines(lines: list, threshold: float = 0.8) -> list:
    """Remove redundant lines that occupy nearly identical space with similar text."""
    if not lines: return []
    unique = []
    for l in lines:
        is_dup = False
        l_text = l["text"].strip().lower()
        l_box = (l["x0"], l["top"], l["x1"], l["bottom"])
        for u in unique:
            u_text = u["text"].strip().lower()
            u_box = (u["x0"], u["top"], u["x1"], u["bottom"])
            
            box_match = (abs(l_box[0] - u_box[0]) < 5 and abs(l_box[1] - u_box[1]) < 5)
            text_match = (l_text == u_text or l_text in u_text or u_text in l_text)
            if box_match and text_match:
                is_dup = True
                break
        if not is_dup:
            unique.append(l)
    return unique



def _is_label_value_line(label: str) -> bool:
    """Return True only when the left side of a colon looks like a plain label.
    Guards against splitting URLs, timestamps, ratios, codes, etc."""
    label = label.strip()
    if len(label) > 40:
        return False
    if any(ch.isdigit() for ch in label):
        return False
    if '/' in label or '..' in label:
        return False
    return True



def _fix_punctuation(text: str) -> str:
    """Remove spaces that Google Translate inserts before punctuation marks."""
    
    text = re.sub(r'\s+([,\.;:!?\)\]}\u0964\u0965])', r'\1', text)
    
    text = re.sub(r'([\(\[{])\s+', r'\1', text)
   
    if any(0x0900 <= ord(c) <= 0x097F for c in text):
        text = re.sub(
            r',\s*(\u0939\u0948|\u0939\u0948\u0902|\u0925\u093e|\u0925\u0947|\u0925\u0940'
            r'|\u0939\u094b|\u0939\u094b\u0917\u093e|\u0939\u094b\u0917\u0940'
            r'|\u0939\u094b\u0902\u0917\u0947|\u0939\u0942\u0901|\u0939\u0942\u0902)(?=[\s\.,\u0964\u0965\?!]|$)',
            r' \1', text
        )
    return text.strip()


def _get_unicode_font_id(pdf_out, fs: float) -> str:
    """Ensure a Unicode/Devanagari font is registered and active; return its font ID."""
    deva_ttf = SCRIPT_FONTS.get("Devanagari")
    if deva_ttf and deva_ttf.endswith(".ttf"):
        font_id = os.path.basename(deva_ttf).split(".")[0]
        if font_id.lower() not in pdf_out.fonts:
            pdf_out.add_font(font_id, fname=deva_ttf)
        pdf_out.set_font(font_id, size=fs)
        return font_id
    
    fallback_ttf = SCRIPT_FONTS.get("Regular")
    if fallback_ttf and fallback_ttf.endswith(".ttf"):
        font_id = os.path.basename(fallback_ttf).split(".")[0]
        if font_id.lower() not in pdf_out.fonts:
            pdf_out.add_font(font_id, fname=fallback_ttf)
        pdf_out.set_font(font_id, size=fs)
        return font_id
    pdf_out.set_font("helvetica", size=fs)
    return "helvetica"



def _detect_source_lang(text: str) -> str:
    """Return 'en' if the text is mostly ASCII/Latin, else 'auto'."""
    if not text:
        return 'auto'
    alpha_chars = [c for c in text if c.isalpha()]
    if not alpha_chars:
        return 'auto'
    ascii_alpha = sum(1 for c in alpha_chars if ord(c) < 128)
    return 'en' if (ascii_alpha / len(alpha_chars)) > 0.75 else 'auto'



_GLOSSARY: dict = {
    "hi": {
        "बकाई": "एआई",
        "विकाई": "एआई",
        "छकाई": "एआई",
        "एकाई": "एआई",   
        "आकाई": "एआई",
        "artificial intelligence": "कृत्रिम बुद्धिमत्ता",
        "machine learning": "मशीन लर्निंग",
        "deep learning": "डीप लर्निंग",
        "natural language processing": "प्राकृतिक भाषा प्रसंस्करण",
        "neural network": "न्यूरल नेटवर्क",
    },
    "mr": {
        "बकाई": "एआई",
        "विकाई": "एआई",
        "एकाई": "एआई",
        "artificial intelligence": "कृत्रिम बुद्धिमत्ता",
        "machine learning": "मशीन लर्निंग",
    },
}


def _apply_glossary(text: str, lang: str) -> str:
    """Case-insensitive replacement of known bad translations."""
    glossary = _GLOSSARY.get(lang, {})
    for wrong, correct in glossary.items():
        
        text = re.sub(re.escape(wrong), correct, text, flags=re.IGNORECASE)
    return text


def _is_non_latin(text: str) -> bool:
    """Return True if the text contains any non-ASCII alphabetic character."""
    return any(c.isalpha() and ord(c) > 127 for c in text)



def _translate_batch(texts: list, source_lang: str, target_lang: str) -> list:
    """Translate multiple strings in one API call using a numbered-list format.
    Falls back to individual calls if batching fails or the payload is too large."""
    import time
    if not texts:
        return texts

    numbered = "\n".join(f"{i + 1}) {t.strip()}" for i, t in enumerate(texts))
    if len(numbered) > 4500:        
        return [_fix_punctuation(_translate_with_retry(t, target_lang)) for t in texts]

    try:
        result = GoogleTranslator(source=source_lang, target=target_lang).translate(numbered)
        if result:
            parsed: dict = {}
            for line in result.strip().splitlines():
                m = re.match(r'^(\d+)[.)\s]\s*(.*)', line.strip())
                if m:
                    parsed[int(m.group(1)) - 1] = m.group(2).strip()
            if len(parsed) == len(texts):
                return [_fix_punctuation(parsed.get(i, texts[i])) for i in range(len(texts))]
    except Exception as e:
        print(f"[batch_translate] failed: {e}")

    # Fallback: translateone by one
    return [_fix_punctuation(_translate_with_retry(t, target_lang)) for t in texts]


# ── FIX 6: Column detector ─────────────────────────────────────────────────────
def _detect_columns(lines: list, page_width: float) -> int:
    """Return 1 or 2 — how many text columns the page likely has."""
    if len(lines) < 6:
        return 1
    x0s = [l["x0"] for l in lines]
    # Count lines whose x0 sits in the right half of the page
    right_count = sum(1 for x in x0s if x > page_width * 0.45)
    # If >25 % of lines start on the right half, assume 2-column layout
    if right_count > len(x0s) * 0.25:
        return 2
    return 1


def process_translation(input_path: str, language: str, task_id: str, file_key: str):
    import pdfplumber
    import fitz
    import os
    import time

    from database import SessionLocal
    db = SessionLocal()

    try:
        output_path = f"translated_{task_id}.pdf"
        image_temp_files = []

        # Open fitz doc alongside pdfplumber for image OCR (Fix 5)
        fitz_doc = fitz.open(input_path)

        with pdfplumber.open(input_path) as pdf:
            total_pages = len(pdf.pages)
            first_page  = pdf.pages[0]

            # FIX 3: Detect source language once from first page
            sample_text = (first_page.extract_text() or "")[:600]
            source_lang = _detect_source_lang(sample_text)
            print(f"[translate] source_lang detected: {source_lang}")

            # FIX 2: No default format — set per-page in add_page()
            pdf_out = FPDF(unit="pt")
            pdf_out.set_text_shaping(True)
            # FIX 2: Disable auto page-break so FPDF never inserts phantom pages
            pdf_out.set_auto_page_break(False)

            for page_index, page in enumerate(pdf.pages):
                current_page_num = page_index + 1

                # FIX 2: Always pass actual page size — no special-casing for page 0
                pdf_out.add_page(format=(page.width, page.height))

                fitz_page = fitz_doc[page_index]

                # ── PROCESS IMAGES (FIX 5) ──────────────────────────────────────
                processed_rects = [] # list of (x0, top, x1, bottom) for skipping main text loop
                
                for img in page.images:
                    try:
                        bbox = (img["x0"], img["top"], img["x1"], img["bottom"])
                        cropped_page = page.within_bbox(bbox)
                        img_obj = cropped_page.to_image(resolution=200)

                        temp_img = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                        img_obj.save(temp_img.name, format="PNG")
                        temp_img.close()
                        image_temp_files.append(temp_img.name)

                        # Fix 2: Wipe the full image bbox first so any PDF text-layer
                        # sitting underneath (English chart annotations) is hidden
                        pdf_out.set_fill_color(255, 255, 255)
                        pdf_out.rect(img["x0"], img["top"], img["width"], img["height"], style="F")

                        # Draw the rasterised image
                        pdf_out.image(temp_img.name,
                                      x=img["x0"], y=img["top"],
                                      w=img["width"], h=img["height"])

                        # FIX 5A: Extract vector text from the image region via fitz
                        # This catches labels like "Time", "Impact" often found in charts
                        clip = fitz.Rect(img["x0"], img["top"],
                                         img["x1"], img["bottom"])
                        vec_blocks = fitz_page.get_text("blocks", clip=clip)

                        found_vec_text = False
                        for block in (vec_blocks or []):
                            btext = (block[4] if len(block) > 4 else "").strip()
                            if not btext or len(btext) < 2:
                                continue
                            
                            found_vec_text = True
                            bx0, by0, bx1, by1 = block[0], block[1], block[2], block[3]
                            
                            # Add to processed_rects to avoid drawing this text again in the main loop
                            processed_rects.append((bx0 - 2, by0 - 2, bx1 + 2, by1 + 2))
                            
                            t_btext = _fix_punctuation(_translate_with_retry(btext, language))
                            
                            # Fix 2: Expand white rect ±3 pt so no edge pixel of English text bleeds through
                            pdf_out.set_fill_color(255, 255, 255)
                            pdf_out.rect(bx0 - 3, by0 - 3, (bx1 - bx0) + 6, (by1 - by0) + 6, style="F")
                            # Use smaller font for chart labels; pass lang_code for Fix 1
                            draw_mixed_text(pdf_out, t_btext, bx0, by1, 8, lang_code=language)

                        # FIX 5B: OCR fallback for raster images without vector text
                        if not found_vec_text:
                            try:
                                import pytesseract
                                # Run tesseract to get bounding boxes and text
                                ocr_data = pytesseract.image_to_data(temp_img.name, output_type=pytesseract.Output.DICT)
                                for i in range(len(ocr_data["text"])):
                                    txt = ocr_data["text"][i].strip()
                                    if len(txt) > 2 and int(ocr_data["conf"][i]) > 50:
                                        # Scale coordinates from raster to PDF
                                        scale_x = img["width"] / img_obj.width
                                        scale_y = img["height"] / img_obj.height
                                        ox0 = img["x0"] + ocr_data["left"][i] * scale_x
                                        oy0 = img["top"] + ocr_data["top"][i] * scale_y
                                        ox1 = ox0 + ocr_data["width"][i] * scale_x
                                        oy1 = oy0 + ocr_data["height"][i] * scale_y
                                        
                                        t_txt = _fix_punctuation(_translate_with_retry(txt, language))
                                        
                                        # White out and overdraw
                                        pdf_out.set_fill_color(255, 255, 255)
                                        pdf_out.rect(ox0 - 2, oy0 - 2, (ox1 - ox0) + 4, (oy1 - oy0) + 4, style="F")
                                        draw_mixed_text(pdf_out, t_txt, ox0, oy1, 8, lang_code=language)
                            except Exception as ocr_err:
                                print(f"[image OCR] skipped: {ocr_err}")

                    except Exception as img_err:
                        print(f"[image] page {current_page_num}: {img_err}")


                # ── PROCESS TEXT ─────────────────────────────────────────────
                # FIX 1: Primary extraction + Word fallback
                raw_lines = page.extract_text_lines(layout=True, x_tolerance=2)
                
                # Word-level fallback logic to catch missed content
                line_rects = [(l["x0"], l["top"], l["x1"], l["bottom"]) for l in raw_lines]
                uncovered = []
                for word in (page.extract_words() or []):
                    wx0, wtop, wx1, wbot = word["x0"], word["top"], word["x1"], word["bottom"]
                    if not any(lx0-2 <= wx0 and lx1+2 >= wx1 and ltop-2 <= wtop and lbot+2 >= wbot for lx0, ltop, lx1, lbot in line_rects):
                        uncovered.append(word)
                if uncovered:
                    uncovered.sort(key=lambda w: (round(w["top"]/5)*5, w["x0"]))
                    groups, cur_grp = [], [uncovered[0]]
                    for w in uncovered[1:]:
                        if abs(w["top"] - cur_grp[-1]["top"]) < 6: cur_grp.append(w)
                        else: groups.append(cur_grp); cur_grp = [w]
                    groups.append(cur_grp)
                    for grp in groups:
                        raw_lines.append({"text": " ".join(w["text"] for w in grp), "x0": min(w["x0"] for w in grp), "x1": max(w["x1"] for w in grp), "top": min(w["top"] for w in grp), "bottom": max(w["bottom"] for w in grp)})

                # FIX 1: Deduplicate and filter out regions already processed in images/charts
                text_lines = []
                for line in _deduplicate_lines(raw_lines):
                    # Check if this line overlaps with any rect already processed (e.g. chart labels)
                    lx0, ltop, lx1, lbot = line["x0"], line["top"], line["x1"], line["bottom"]
                    is_processed = any(px0 <= lx0+2 and px1 >= lx1-2 and ptop <= ltop+2 and pbot >= lbot-2 for px0, ptop, px1, pbot in processed_rects)
                    if not is_processed:
                        text_lines.append(line)

                text_lines.sort(key=lambda l: (l["top"], l["x0"]))


                # Annotate lines with font size and column id
                for line in text_lines:
                    chars_on = [
                        c for c in page.chars
                        if c["top"] >= line["top"] - 1
                        and c["bottom"] <= line["bottom"] + 1
                    ]
                    line["_fs"] = (
                        sum(c["size"] for c in chars_on) / len(chars_on)
                        if chars_on else 12
                    )

                # FIX 6: Detect columns; build per-column Y cursors
                num_cols = _detect_columns(text_lines, page.width)
                col_y: dict = {}  # col_id -> current baseline y

                def _col_id(x0: float) -> int:
                    return 0 if (num_cols == 1 or x0 < page.width * 0.48) else 1

                for line in text_lines:
                    line["_col"] = _col_id(line["x0"])

                # FIX 3 / FIX 6: Build semantic batches based on vertical proximity and column
                # Fix 4: Track translated blocks already drawn on this page to prevent duplication
                seen_blocks: set = set()
                seen_source_blocks: set = set()

                batches = []
                cur_batch = []
                last_y = -100
                last_col = -1
                
                for line in text_lines:
                    t = line["text"].strip()
                    if not t: continue
                    
                    fs = line["_fs"]
                    col = line["_col"]
                    y = line["bottom"]
                    
                    # Start new batch if column changes or there is a large vertical jump (new paragraph)
                    if col != last_col or abs(y - last_y) > fs * 2.5 or len(cur_batch) >= 10:
                        if cur_batch: batches.append(cur_batch)
                        cur_batch = [line]
                    else:
                        cur_batch.append(line)
                    
                    last_y = y
                    last_col = col
                if cur_batch: batches.append(cur_batch)

                # Translate each batch then draw
                for batch in batches:
                    # Semantic Joining: Join lines into a single string for better context
                    raw_text = " ".join(l["text"].strip() for l in batch)
                    _source_key = raw_text.lower()
                    
                    # Fix 4 (Round 2): Overlapping shadow text forms duplicate extracted English lines
                    # If the source string is >85% identical to an already seen source block on this page, it's a shadow duplicate
                    from difflib import SequenceMatcher
                    found_dup = False
                    if len(_source_key) > 5:
                        for seen_src in seen_source_blocks:
                            if SequenceMatcher(None, _source_key, seen_src).ratio() > 0.85:
                                found_dup = True
                                break
                    
                    if found_dup:
                        continue
                    if len(_source_key) > 5:
                        seen_source_blocks.add(_source_key)

                    if ":" in raw_text and len(batch) == 1:
                        parts = raw_text.split(":", 1)
                        lbl, val = parts[0].strip(), parts[1].strip()
                        if _is_label_value_line(lbl):
                            t_lbl = _translate_with_retry(lbl, language)
                            t_val = _translate_with_retry(val, language) if val else ""
                            translated_block = _fix_punctuation(f"{t_lbl} : {t_val}")
                        else:
                            translated_block = _fix_punctuation(_translate_with_retry(raw_text, language))
                    else:
                        translated_block = _fix_punctuation(_translate_with_retry(raw_text, language))

                    # Apply glossary to fix badly translated AI/tech terms
                    translated_block = _apply_glossary(translated_block, language)

                    # Fix 4: Skip this batch if we already drew the same translated text on this page
                    _block_key = " ".join(translated_block.split())
                    if len(_block_key) >= 4:
                        if _block_key in seen_blocks:
                            continue
                        seen_blocks.add(_block_key)

                    first_line = batch[0]
                    last_line  = batch[-1]
                    fs  = first_line["_fs"]
                    col = first_line["_col"]

                    if col not in col_y or col_y[col] < first_line["bottom"] - fs * 2:
                        col_y[col] = first_line["bottom"]

                    draw_y = min(col_y[col], page.height - 4)

                    # ── SHADOW TEXT FIX ──────────────────────────────────────
                    # White out the original text bounding box for every line
                    # in the batch so translated text is not drawn on top.
                    pdf_out.set_fill_color(255, 255, 255)
                    for orig_line in batch:
                        bx0  = orig_line["x0"]
                        btop = orig_line["top"]
                        bw   = orig_line["x1"] - bx0
                        bh   = orig_line["bottom"] - btop
                        if bw > 0 and bh > 0:
                            pdf_out.rect(bx0, btop, bw, bh + 1, style="F")

                    # ── WIDTH MEASUREMENT using Unicode font ──────────────────
                    col_width = (page.width * 0.45) if num_cols == 2 else (page.width - first_line["x0"] - 20)

                    # Always use a Unicode-capable font for measurement to avoid crashes
                    if _is_non_latin(translated_block):
                        _get_unicode_font_id(pdf_out, fs)
                    else:
                        pdf_out.set_font("helvetica", size=fs)

                    words = translated_block.split()
                    line_chunk = ""
                    for word in words:
                        test_str = (line_chunk + " " + word).strip()
                        try:
                            w = pdf_out.get_string_width(test_str)
                        except (UnicodeEncodeError, FPDFUnicodeEncodingException):
                            w = len(test_str) * (fs * 0.5)

                        if w < col_width:
                            line_chunk = test_str
                        else:
                            if line_chunk:
                                # Fix 5 (Round 2): blunt force comma cleanup
                                line_chunk = line_chunk.replace(", है", " है").replace(", हैं", " हैं").replace(", था", " था").replace(", थे", " थे").replace(", थी", " थी")
                                # Fix 1: pass lang_code so HarfBuzz uses the correct script tag
                                draw_mixed_text(pdf_out, line_chunk, first_line["x0"], draw_y, fs,
                                                lang_code=language)
                                draw_y += fs * 1.3
                            line_chunk = word
                    if line_chunk:
                        # Fix 5 (Round 2): blunt force comma cleanup
                        line_chunk = line_chunk.replace(", है", " है").replace(", हैं", " हैं").replace(", था", " था").replace(", थे", " थे").replace(", थी", " थी")
                        draw_mixed_text(pdf_out, line_chunk, first_line["x0"], draw_y, fs,
                                        lang_code=language)
                        draw_y += fs * 1.3

                    col_y[col] = draw_y



                # Update progress
                progress_store[task_id] = int(
                    (current_page_num / total_pages) * 90)
                print(f"[translate] page {current_page_num}/{total_pages} done")

        fitz_doc.close()

        # Finalise PDF
        pdf_out.output(output_path)

        # Cleanup temp images
        for tmp in image_temp_files:
            try:
                os.remove(tmp)
            except Exception:
                pass

        # Upload to S3
        progress_store[task_id] = 95
        s3.upload_file(
            output_path, BUCKET_NAME, file_key,
            ExtraArgs={"ContentType": "application/pdf"},
        )

        try:
            os.remove(output_path)
            os.remove(input_path)
        except Exception:
            pass

        progress_store[task_id] = 100
        print("[translate] Translation complete!")

        task_record = db.query(models.TranslationTask).filter(
            models.TranslationTask.id == task_id).first()
        if task_record:
            task_record.status = "completed"
            file_url = s3.generate_presigned_url(
                "get_object",
                Params={"Bucket": BUCKET_NAME, "Key": file_key},
                ExpiresIn=3600 * 24 * 7,   # 7 days
            )
            task_record.download_url = file_url
            db.commit()

    except Exception as e:
        import traceback
        print(f"[translate] Error: {e}")
        traceback.print_exc()
        progress_store[task_id] = -1
        task_record = db.query(models.TranslationTask).filter(
            models.TranslationTask.id == task_id).first()
        if task_record:
            task_record.status = "failed"
            db.commit()
    finally:
        db.close()

# ---------------------------
# DOCX Translation (Background)
# ---------------------------

def process_docx_translation(input_path: str, language: str, task_id: str, file_key: str):
    """Translate a DOCX file paragraph-by-paragraph and upload to S3."""
    import os
    from docx import Document
    from docx.shared import Pt
    from database import SessionLocal
    db = SessionLocal()
    try:
        doc = Document(input_path)
        total_paras = len(doc.paragraphs)
        output_path = f"translated_{task_id}.docx"

        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                translated = _translate_with_retry(para.text, language)
                translated = _fix_punctuation(translated)
                translated = _apply_glossary(translated, language)
                # Preserve runs structure – replace first run, clear rest
                if para.runs:
                    para.runs[0].text = translated
                    for run in para.runs[1:]:
                        run.text = ""
                else:
                    para.text = translated
            progress_store[task_id] = int(((i + 1) / max(total_paras, 1)) * 88)

        doc.save(output_path)

        # Upload translated DOCX to S3
        progress_store[task_id] = 95
        s3.upload_file(
            output_path, BUCKET_NAME, file_key,
            ExtraArgs={"ContentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
        )

        try:
            os.remove(output_path)
            os.remove(input_path)
        except Exception:
            pass

        progress_store[task_id] = 100
        print("[docx-translate] Translation complete!")

        task_record = db.query(models.TranslationTask).filter(
            models.TranslationTask.id == task_id).first()
        if task_record:
            task_record.status = "completed"
            file_url = s3.generate_presigned_url(
                "get_object",
                Params={"Bucket": BUCKET_NAME, "Key": file_key},
                ExpiresIn=3600 * 24 * 7,
            )
            task_record.download_url = file_url
            db.commit()

    except Exception as e:
        import traceback
        print(f"[docx-translate] Error: {e}")
        traceback.print_exc()
        progress_store[task_id] = -1
        task_record = db.query(models.TranslationTask).filter(
            models.TranslationTask.id == task_id).first()
        if task_record:
            task_record.status = "failed"
            db.commit()
    finally:
        db.close()


# ---------------------------
# Translate PDF/DOCX API Upload
# ---------------------------

ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "application/octet-stream",
}

@app.post("/translate")
async def translate_pdf(
    request: Request,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    language: str = Form(...),
    current_user: models.User = Depends(auth.get_current_user),
    db: Session = Depends(get_db)
):
    filename_lower = (file.filename or "").lower()
    is_pdf  = filename_lower.endswith(".pdf")
    is_docx = filename_lower.endswith(".docx")

    # Strict file-type gate – only PDF and DOCX allowed
    if not is_pdf and not is_docx:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Please upload a PDF (.pdf) or Word document (.docx) only."
        )

    contents = await file.read()

    # Validate file size (max 20 MB)
    MAX_SIZE = 20 * 1024 * 1024
    if len(contents) > MAX_SIZE:
        raise HTTPException(status_code=400, detail="File too large. Maximum allowed size is 20 MB.")

    # Magic-byte validation
    if is_pdf and not contents.startswith(b"%PDF"):
        raise HTTPException(status_code=400, detail="Invalid PDF file. The file content does not match PDF format.")
    if is_docx and not contents.startswith(b"PK"):  # DOCX/ZIP magic bytes
        raise HTTPException(status_code=400, detail="Invalid DOCX file. The file content does not match Word format.")

    suffix  = ".pdf" if is_pdf else ".docx"
    file_key_suffix = ".pdf" if is_pdf else ".docx"

    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as temp:
        temp.write(contents)
        input_path = temp.name

    task_id  = str(uuid.uuid4())
    file_key = f"translated/{task_id}{file_key_suffix}"

    # Generate initial pre-signed download URL
    file_url = s3.generate_presigned_url(
        "get_object",
        Params={"Bucket": BUCKET_NAME, "Key": file_key},
        ExpiresIn=3600
    )

    # Save initial task record to DB
    new_task = models.TranslationTask(
        id=task_id,
        user_id=current_user.id,
        original_filename=file.filename,
        target_language=language,
        status="processing"
    )
    db.add(new_task)
    db.commit()

    # Dispatch to correct background worker
    if is_docx:
        background_tasks.add_task(
            process_docx_translation,
            input_path=input_path,
            language=language,
            task_id=task_id,
            file_key=file_key
        )
    else:
        background_tasks.add_task(
            process_translation,
            input_path=input_path,
            language=language,
            task_id=task_id,
            file_key=file_key
        )

    return {
        "task_id": task_id,
        "download_url": file_url
    }

# ── Serve frontend static files (at the end to avoid shadowing API) ───────────
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, RedirectResponse

# Serve specific HTML files from root for easier deployment
@app.get("/", include_in_schema=False)
async def serve_root():
    if os.path.exists("login.html"):
        return FileResponse("login.html")
    return {"message": "Worldocs API Running"}

@app.get("/index", include_in_schema=False)
async def index_redirect():
    return RedirectResponse(url="/index.html")

@app.get("/login", include_in_schema=False)
async def login_redirect():
    return RedirectResponse(url="/login.html")

@app.get("/dashboard", include_in_schema=False)
async def dashboard_redirect():
    return RedirectResponse(url="/dashboard.html")

@app.get("/{filename}", include_in_schema=False)
async def serve_static_root(filename: str):
    # List of allowed static file extensions
    allowed_extensions = {".css", ".js", ".ttf", ".png", ".jpg", ".jpeg", ".svg", ".pdf", ".ico"}
    _, ext = os.path.splitext(filename)
    
    if ext.lower() in allowed_extensions:
        if os.path.exists(filename):
            return FileResponse(filename)
            
    # Also handle .html if requested explicitly or without extension
    if filename.endswith(".html") or not ext:
        html_path = filename if filename.endswith(".html") else f"{filename}.html"
        if os.path.exists(html_path):
            return FileResponse(html_path)
        
    raise HTTPException(status_code=404)

# If the user has a 'frontend' directory, mount it.
if os.path.isdir("frontend"):
    app.mount("/frontend", StaticFiles(directory="frontend"), name="frontend")